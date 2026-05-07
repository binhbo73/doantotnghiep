"""
Document Parser (OPTIMIZED)
============================
Extract text from documents (PDF, DOCX, TXT, Markdown, Excel)

OPTIMIZATION:
- DOCX: python-docx (5-10x faster than docling)
- PDF: docling (quality > speed)
- TXT/MD: Standard file read
- XLSX: openpyxl (fast, structured data)
- CACHE: Redis for parsed results (TTL 7 days)

Performance benchmark (30-page Word):
- docling: ~30-45 seconds
- python-docx: ~2-5 seconds  ✓ 10x faster

Features:
- Fast DOCX parsing via python-docx
- Fast XLSX parsing via openpyxl
- Quality PDF parsing via docling
- Redis caching to avoid re-parsing
- Metadata extraction
- Error handling

Configuration (from settings.py):
    DOCLING_MAX_FILE_SIZE_MB = 100
    DOCLING_TIMEOUT = 60

Usage:
    parser = DocumentParser()
    
    # Parse from file (checks cache first)
    text, metadata = parser.parse_file('/path/to/document.xlsx')
    
    # Parse DOCX (fast via python-docx)
    docx_text = parser.parse_docx(file_path)
    
    # Parse Excel (fast via openpyxl)
    xlsx_text = parser.parse_xlsx(file_path)
    
    # Parse PDF (quality via docling)
    pdf_text = parser.parse_pdf(file_path)
"""

import logging
import os
import mimetypes
import hashlib
import re
from typing import Tuple, Dict, Any, Optional
from pathlib import Path
from django.conf import settings
from core.exceptions import DocumentProcessingError

logger = logging.getLogger(__name__)


class DocumentParser:
    """
    Document parser - extracts text from various formats (OPTIMIZED)
    
    Supported formats:
    - PDF (.pdf) → docling (quality)
    - DOCX (.docx, .doc) → python-docx (FAST, 10x)
    - TXT (.txt) → standard read
    - Markdown (.md) → standard read
    
    Cache: Uses Redis to store parsed results
    """
    
    SUPPORTED_TYPES = {
        'application/pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',  # .docx
        'application/msword',  # .doc
        'text/plain',
        'text/markdown',
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',  # .xlsx
        'application/vnd.ms-excel',  # .xls
    }
    
    CACHE_TTL_SECONDS = 7 * 24 * 3600  # 7 days
    
    def __init__(
        self,
        max_file_size_mb: int = None,
        timeout: int = None,
        use_cache: bool = True,
    ):
        """
        Initialize parser
        
        Args:
            max_file_size_mb: Max file size in MB (default from settings)
            timeout: Processing timeout in seconds (default 60)
            use_cache: Enable Redis caching (default True)
        """
        self.max_file_size_mb = max_file_size_mb or getattr(
            settings, 'DOCLING_MAX_FILE_SIZE_MB', 100
        )
        self.timeout = timeout or getattr(settings, 'DOCLING_TIMEOUT', 60)
        self.use_cache = use_cache
        
        self.redis_client = None
        if use_cache:
            try:
                import redis
                self.redis_client = redis.StrictRedis(
                    host=getattr(settings, 'REDIS_HOST', 'localhost'),
                    port=getattr(settings, 'REDIS_PORT', 6379),
                    db=getattr(settings, 'REDIS_DB', 1),
                    decode_responses=True,
                )
                # Test connection
                self.redis_client.ping()
                logger.info("DocumentParser cache enabled (Redis)")
            except Exception as e:
                logger.warning(f"Redis cache disabled: {e}")
                self.redis_client = None
        
        logger.info(f"DocumentParser initialized: max_size={self.max_file_size_mb}MB, cache={use_cache}")
    
    # ============================================================================
    # CACHE METHODS
    # ============================================================================
    
    def _get_cache_key(self, file_path: str) -> str:
        """Generate cache key from file content hash"""
        if not os.path.exists(file_path):
            return None
        
        # Hash based on file size + modification time
        stat = os.stat(file_path)
        key_input = f"{file_path}:{stat.st_size}:{stat.st_mtime}"
        key_hash = hashlib.md5(key_input.encode()).hexdigest()
        return f"doc_parse:{key_hash}"
    
    def _get_cached_result(self, file_path: str) -> Optional[Tuple[str, Dict]]:
        """Get cached parse result if available"""
        if not self.redis_client:
            return None
        
        cache_key = self._get_cache_key(file_path)
        if not cache_key:
            return None
        
        try:
            cached = self.redis_client.get(cache_key)
            if cached:
                import json
                data = json.loads(cached)
                logger.debug(f"Cache hit: {cache_key}")
                return data['text'], data['metadata']
        except Exception as e:
            logger.warning(f"Cache read error: {e}")
        
        return None
    
    def _set_cached_result(self, file_path: str, text: str, metadata: Dict):
        """Store parse result in cache"""
        if not self.redis_client:
            return
        
        cache_key = self._get_cache_key(file_path)
        if not cache_key:
            return
        
        try:
            import json
            data = {
                'text': text,
                'metadata': metadata,
            }
            self.redis_client.setex(
                cache_key,
                self.CACHE_TTL_SECONDS,
                json.dumps(data),
            )
            logger.debug(f"Cached result: {cache_key}")
        except Exception as e:
            logger.warning(f"Cache write error: {e}")
    
    # ============================================================================
    # MAIN PARSING METHOD
    # ============================================================================
    
    def parse_file(self, file_path: str, file_type: Optional[str] = None) -> Tuple[str, Dict[str, Any]]:
        """
        Parse document file and extract text + metadata
        
        Checks cache FIRST before parsing.
        
        Args:
            file_path: Path to document file
            file_type: Optional MIME type. If not provided, will guess from extension.
        
        Returns:
            (text, metadata) tuple
            metadata: {
                'title': str,
                'pages': int,
                'word_count': int,
                'file_type': str,
                'language': str,
                'from_cache': bool,
            }
        
        Raises:
            DocumentProcessingError: If parsing fails
        """
        try:
            # 1. Check cache first
            cached = self._get_cached_result(file_path)
            if cached:
                text, metadata = cached
                metadata['from_cache'] = True
                return text, metadata
            
            # 2. Validate file exists
            if not os.path.exists(file_path):
                raise DocumentProcessingError(f"File not found: {file_path}")
            
            # 3. Validate file size
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            if file_size_mb > self.max_file_size_mb:
                raise DocumentProcessingError(
                    f"File too large: {file_size_mb:.1f}MB > {self.max_file_size_mb}MB"
                )
            
            # 4. Get file type if not provided
            if not file_type:
                file_type = mimetypes.guess_type(file_path)[0] or 'unknown'
            
            # 5. Parse based on type
            if file_type == 'application/pdf':
                text, page_count = self.parse_pdf(file_path)
            elif file_type in ('application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword'):
                text = self.parse_docx(file_path)  # ⚡ FAST via python-docx
                page_count = None
            elif file_type in ('application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 'application/vnd.ms-excel'):
                text = self.parse_excel(file_path, file_type=file_type)
                page_count = None
            elif file_type in ('text/plain', 'text/markdown'):
                text = self.parse_text(file_path)
                page_count = None
            else:
                raise DocumentProcessingError(f"Unsupported file type: {file_type}")

            # 6. Normalize extracted text to improve chunk quality for RAG
            text = self._normalize_extracted_text(text)
            
            # 7. Extract metadata
            metadata = self._extract_metadata(text, file_path, file_type, page_count=page_count)
            metadata['from_cache'] = False
            
            # 8. Cache result
            self._set_cached_result(file_path, text, metadata)
            
            logger.info(
                f"Parsed {os.path.basename(file_path)}: {len(text)} chars, "
                f"{metadata['word_count']} words, {metadata['pages']} pages"
            )
            
            return text, metadata
        
        except Exception as e:
            logger.error(f"Error parsing file {file_path}: {str(e)}", exc_info=True)
            raise DocumentProcessingError(f"Failed to parse file: {str(e)}")
    
    # ============================================================================
    # FORMAT-SPECIFIC PARSERS
    # ============================================================================
    
    def parse_pdf(self, file_path: str) -> Tuple[str, int]:
        """
        Parse PDF file and extract text (via opendataloader-pdf - #1 benchmarks)
        
        Uses opendataloader-pdf library for superior PDF parsing (better than docling)
        - #1 overall accuracy (0.907 vs docling 0.882)
        - Better table extraction (0.928 TEDS vs docling 0.887)
        - Bounding boxes for citations
        - Hybrid AI mode for complex pages
        
        Args:
            file_path: Path to PDF file
        
        Returns:
            Extracted text (Markdown format for RAG)
        
        Raises:
            DocumentProcessingError: If parsing fails
        """
        import tempfile
        import shutil
        
        try:
            from pypdf import PdfReader
            import opendataloader_pdf
            
            logger.debug(f"Parsing PDF (opendataloader-pdf): {os.path.basename(file_path)}")

            page_count = 0
            try:
                with open(file_path, 'rb') as pdf_file:
                    page_count = len(PdfReader(pdf_file).pages)
            except Exception as page_err:
                logger.warning(f"Could not read PDF page count for {file_path}: {page_err}")
            
            # Create temp directory for output
            with tempfile.TemporaryDirectory() as temp_dir:
                # Convert PDF using opendataloader-pdf
                # Use markdown-with-images + embedded images so extracted image data is preserved
                opendataloader_pdf.convert(
                    input_path=[file_path],
                    output_dir=temp_dir,
                    format="markdown-with-images",
                    image_output="embedded",
                    image_format="png"
                )
                
                # Find the output markdown file
                # opendataloader-pdf names output as {original_name}.md
                base_name = os.path.splitext(os.path.basename(file_path))[0]
                md_file = os.path.join(temp_dir, f"{base_name}.md")
                
                if not os.path.exists(md_file):
                    raise DocumentProcessingError(f"Output markdown file not found: {md_file}")
                
                # Read the extracted text
                with open(md_file, 'r', encoding='utf-8') as f:
                    text = f.read()
                
                logger.debug(f"PDF parsing completed: {len(text)} chars")
                return text, page_count or 1
        
        except ImportError as e:
            raise DocumentProcessingError(
                f"opendataloader-pdf dependency missing: {str(e)}. "
                "Install with: pip install opendataloader-pdf"
            )
        except Exception as e:
            logger.error(f"PDF parsing error: {str(e)}", exc_info=True)
            raise DocumentProcessingError(f"Failed to parse PDF: {str(e)}")
    
    def parse_docx(self, file_path: str) -> str:
        """
        Parse DOCX file FAST via python-docx (⚡ 10x faster than docling)
        
        Args:
            file_path: Path to DOCX file
        
        Returns:
            Extracted text
        """
        try:
            from docx import Document
            
            logger.debug(f"Parsing DOCX (python-docx): {os.path.basename(file_path)}")
            
            # Load document
            doc = Document(file_path)
            
            # Extract text from paragraphs
            text_parts = []
            
            # Add all paragraph text
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Add text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    if any(row_data):
                        text_parts.append(' | '.join(row_data))
            
            text = '\n'.join(text_parts)
            
            logger.debug(f"DOCX parsing completed: {len(text)} chars")
            return text
        
        except ImportError:
            raise DocumentProcessingError("python-docx library not installed. Install: pip install python-docx")
        except Exception as e:
            logger.error(f"DOCX parsing error: {str(e)}")
            raise DocumentProcessingError(f"Failed to parse DOCX: {str(e)}")
    
    def parse_text(self, file_path: str) -> str:
        """
        Parse plain text / markdown file
        
        Args:
            file_path: Path to text file
        
        Returns:
            File content
        """
        try:
            logger.debug(f"Parsing text file: {os.path.basename(file_path)}")
            
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            logger.debug(f"Text parsing completed: {len(text)} chars")
            return text
        
        except Exception as e:
            logger.error(f"Text parsing error: {str(e)}")
            raise DocumentProcessingError(f"Failed to parse text file: {str(e)}")
    
    def parse_xlsx(self, file_path: str) -> str:
        """
        Parse Excel file (.xlsx) and extract all data
        
        Strategy: Extract all sheets with structure preserved
        - Sheet name as heading
        - Headers as bold
        - Data rows tab-separated
        - Between sheets: separator line
        
        This format is optimal for chunking + embedding while maintaining context.
        
        Args:
            file_path: Path to Excel file
        
        Returns:
            Extracted text with preserved structure
        """
        try:
            from openpyxl import load_workbook
            
            logger.debug(f"Parsing XLSX (openpyxl): {os.path.basename(file_path)}")
            
            workbook = load_workbook(file_path, data_only=True)
            text_parts = []
            
            for sheet_name in workbook.sheetnames:
                worksheet = workbook[sheet_name]
                
                # Add sheet name as heading
                text_parts.append(f"=== SHEET: {sheet_name} ===")
                
                sheet_data = []
                for row in worksheet.iter_rows(values_only=True):
                    # Convert None to empty string and all values to string
                    row_data = [str(cell) if cell is not None else '' for cell in row]
                    # Remove trailing empty cells
                    while row_data and row_data[-1] == '':
                        row_data.pop()
                    
                    if row_data:  # Only add non-empty rows
                        sheet_data.append('\t'.join(row_data))
                
                if sheet_data:
                    text_parts.extend(sheet_data)
                
                # Add separator between sheets
                text_parts.append('\n' + '='*50 + '\n')
            
            text = '\n'.join(text_parts)
            
            logger.debug(f"XLSX parsing completed: {len(text)} chars from {len(workbook.sheetnames)} sheets")
            return text
        
        except ImportError:
            raise DocumentProcessingError("openpyxl library not installed. Install: pip install openpyxl")
        except Exception as e:
            logger.error(f"XLSX parsing error: {str(e)}", exc_info=True)
            raise DocumentProcessingError(f"Failed to parse XLSX: {str(e)}")

    def parse_excel(self, file_path: str, file_type: Optional[str] = None) -> str:
        """
        Parse Excel files both .xlsx and .xls.

        Delegates to the correct parser implementation based on file extension or MIME type.
        """
        if (file_type == 'application/vnd.ms-excel' or file_path.lower().endswith('.xls')):
            return self.parse_xls(file_path)
        return self.parse_xlsx(file_path)

    def _normalize_extracted_text(self, text: str) -> str:
        """
        Normalize parser output before chunking.

        Why:
        - Doc parsers may emit markdown/html image markers like <!-- image -->
        - Placeholder-only tokens reduce semantic quality for embeddings
        """
        if not text:
            return text

        normalized = text

        # Convert markdown images to readable text so image captions/alt survive.
        normalized = re.sub(
            r'!\[([^\]]*)\]\(([^)]+)\)',
            lambda m: f"[Image: {(m.group(1) or 'embedded image').strip()}]",
            normalized,
            flags=re.IGNORECASE,
        )

        # Convert HTML <img ...> tags to readable text using alt/title when present.
        def _replace_img_tag(match):
            tag = match.group(0)
            alt_match = re.search(r'alt=[\"\']([^\"\']+)[\"\']', tag, flags=re.IGNORECASE)
            title_match = re.search(r'title=[\"\']([^\"\']+)[\"\']', tag, flags=re.IGNORECASE)
            label = None
            if alt_match and alt_match.group(1).strip():
                label = alt_match.group(1).strip()
            elif title_match and title_match.group(1).strip():
                label = title_match.group(1).strip()
            return f"[Image: {label or 'embedded image'}]"

        normalized = re.sub(
            r'<img\b[^>]*>',
            _replace_img_tag,
            normalized,
            flags=re.IGNORECASE,
        )

        # Remove HTML comment placeholders such as <!-- image --> produced by converters.
        normalized = re.sub(
            r'<!--\s*image\s*-->',
            '[Image]',
            normalized,
            flags=re.IGNORECASE,
        )

        # Drop any remaining HTML comments.
        normalized = re.sub(r'<!--.*?-->', ' ', normalized, flags=re.DOTALL)

        # Keep paragraph structure but avoid excessive blank lines/noisy spaces.
        normalized = re.sub(r'[ \t]+', ' ', normalized)
        normalized = re.sub(r'\n{3,}', '\n\n', normalized)

        return normalized.strip()

    def parse_xls(self, file_path: str) -> str:
        """
        Parse legacy Excel .xls files using xlrd.
        """
        try:
            import xlrd

            logger.debug(f"Parsing XLS (xlrd): {os.path.basename(file_path)}")
            workbook = xlrd.open_workbook(file_path)
            text_parts = []

            for sheet in workbook.sheets():
                text_parts.append(f"=== SHEET: {sheet.name} ===")
                sheet_data = []
                for row_index in range(sheet.nrows):
                    row_values = [
                        str(sheet.cell_value(row_index, col)).strip()
                        for col in range(sheet.ncols)
                    ]
                    while row_values and row_values[-1] == '':
                        row_values.pop()

                    if row_values:
                        sheet_data.append('\t'.join(row_values))

                if sheet_data:
                    text_parts.extend(sheet_data)
                text_parts.append('\n' + '=' * 50 + '\n')

            text = '\n'.join(text_parts)
            logger.debug(f"XLS parsing completed: {len(text)} chars from {len(workbook.sheets())} sheets")
            return text

        except ImportError:
            raise DocumentProcessingError("xlrd library not installed. Install: pip install xlrd")
        except Exception as e:
            logger.error(f"XLS parsing error: {str(e)}", exc_info=True)
            raise DocumentProcessingError(f"Failed to parse XLS: {str(e)}")
    
    # ============================================================================
    # METADATA EXTRACTION
    # ============================================================================
    
    def _extract_metadata(
        self,
        text: str,
        file_path: str,
        file_type: str,
        page_count: Optional[int] = None
    ) -> Dict[str, Any]:
        """
        Extract metadata from parsed text
        
        Args:
            text: Extracted text
            file_path: Original file path
            file_type: MIME type
        
        Returns:
            Metadata dict
        """
        try:
            # Count words + lines
            lines = text.split('\n')
            words = text.split()
            
            # Use actual PDF page count when available; otherwise estimate.
            pages = page_count if page_count and page_count > 0 else max(1, len(words) // 300)
            
            # Get filename as potential title
            filename = Path(file_path).stem
            
            metadata = {
                'title': filename,
                'pages': pages,
                'word_count': len(words),
                'char_count': len(text),
                'line_count': len(lines),
                'file_type': file_type,
                'language': 'unknown',
            }
            
            logger.debug(f"Metadata: {pages}p, {len(words)} words")
            return metadata
        
        except Exception as e:
            logger.warning(f"Error extracting metadata: {str(e)}")
            return {
                'title': 'Unknown',
                'pages': 0,
                'word_count': 0,
                'char_count': len(text),
                'line_count': 0,
                'file_type': file_type,
                'language': 'unknown',
            }
    
    # ============================================================================
    # UTILITY METHODS
    # ============================================================================
    
    @staticmethod
    def is_supported_type(file_type: str) -> bool:
        """Check if file type is supported"""
        return file_type in DocumentParser.SUPPORTED_TYPES
    
    @staticmethod
    def get_supported_types() -> set:
        """Get set of supported MIME types"""
        return DocumentParser.SUPPORTED_TYPES.copy()

        try:
            from docling.document_converter import DocumentConverter
            
            logger.debug(f"Parsing DOCX: {file_path}")
            
            converter = DocumentConverter()
            result = converter.convert(file_path)
            
            text = result.document.export_to_markdown()
            
            logger.debug(f"DOCX parsing completed: {len(text)} chars")
            return text
        
        except ImportError:
            raise DocumentProcessingError("docling library not installed")
        except Exception as e:
            logger.error(f"DOCX parsing error: {str(e)}")
            raise DocumentProcessingError(f"Failed to parse DOCX: {str(e)}")
    
    def parse_text(self, file_path: str) -> str:
        """
        Parse plain text / markdown file
        
        Args:
            file_path: Path to text file
        
        Returns:
            File content
        """
        try:
            logger.debug(f"Parsing text file: {file_path}")
            
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            logger.debug(f"Text parsing completed: {len(text)} chars")
            return text
        
        except Exception as e:
            logger.error(f"Text parsing error: {str(e)}")
            raise DocumentProcessingError(f"Failed to parse text file: {str(e)}")
    
    # ============================================================================
    # METADATA EXTRACTION
    # ============================================================================
    
    def _extract_metadata(
        self,
        text: str,
        file_path: str,
        file_type: str,
        page_count: Optional[int] = None
    ) -> Dict[str, Any]:
        """
        Extract metadata from parsed text
        
        Args:
            text: Extracted text
            file_path: Original file path
            file_type: MIME type
        
        Returns:
            Metadata dict
        """
        try:
            # Count words + lines
            lines = text.split('\n')
            words = text.split()
            
            # Use actual PDF page count when available; otherwise estimate.
            pages = page_count if page_count and page_count > 0 else max(1, len(words) // 300)
            
            # Get filename as potential title
            filename = Path(file_path).stem
            
            metadata = {
                'title': filename,
                'pages': pages,
                'word_count': len(words),
                'char_count': len(text),
                'line_count': len(lines),
                'file_type': file_type,
                'language': 'unknown',  # Could add language detection
            }
            
            logger.debug(f"Metadata extracted: {metadata}")
            return metadata
        
        except Exception as e:
            logger.warning(f"Error extracting metadata: {str(e)}")
            return {
                'title': 'Unknown',
                'pages': 0,
                'word_count': 0,
                'char_count': len(text),
                'line_count': 0,
                'file_type': file_type,
                'language': 'unknown',
            }
    
    # ============================================================================
    # UTILITY METHODS
    # ============================================================================
    
    @staticmethod
    def is_supported_type(file_type: str) -> bool:
        """Check if file type is supported"""
        return file_type in DocumentParser.SUPPORTED_TYPES
    
    @staticmethod
    def get_supported_types() -> set:
        """Get set of supported MIME types"""
        return DocumentParser.SUPPORTED_TYPES.copy()
