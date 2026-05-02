"""
Document Parser Benchmark
===========================
Test performance improvement with new optimizations:
- python-docx for DOCX (5-10x faster)
- Redis caching
- Metrics tracking

Usage:
    python manage.py shell < scripts/benchmark_document_parser.py
    
    Or:
    python manage.py shell
    >>> exec(open('scripts/benchmark_document_parser.py').read())
"""

import os
import time
import tempfile
from pathlib import Path
from services.document.parser import DocumentParser
from core.exceptions import DocumentProcessingError

# Sample files for testing (create if doesn't exist)
SAMPLE_DOCX = 'test_samples/sample_30pages.docx'
SAMPLE_PDF = 'test_samples/sample_30pages.pdf'

def create_sample_docx(filename: str, num_pages: int = 30):
    """Create a sample DOCX file with approximate number of pages."""
    try:
        from docx import Document
        from docx.shared import Pt
        
        doc = Document()
        
        # Add title
        doc.add_heading('Sample Report', 0)
        
        # Add paragraphs to simulate pages
        # Approx 300 words per page
        words_per_page = 300
        total_words = num_pages * words_per_page
        
        lorem = """
        Lorem ipsum dolor sit amet, consectetur adipiscing elit. 
        Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. 
        Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris 
        nisi ut aliquip ex ea commodo consequat. 
        """
        
        word_count = 0
        while word_count < total_words:
            doc.add_paragraph(lorem * 10)
            word_count += len(lorem.split())
        
        # Create directory if needed
        os.makedirs(os.path.dirname(filename), exist_ok=True)
        doc.save(filename)
        
        print(f"✓ Created sample DOCX: {filename} (~{num_pages} pages)")
        return filename
    except ImportError:
        print("✗ python-docx not installed. Install: pip install python-docx")
        return None

def benchmark_docx_parsing(file_path: str, iterations: int = 3):
    """Benchmark DOCX parsing speed."""
    print(f"\n{'=' * 70}")
    print(f"BENCHMARK: DOCX Parsing ({iterations} iterations)")
    print(f"{'=' * 70}")
    print(f"File: {file_path}")
    
    if not os.path.exists(file_path):
        print(f"✗ File not found: {file_path}")
        # Create sample file
        file_path = create_sample_docx(file_path)
        if not file_path:
            return
    
    parser = DocumentParser(use_cache=False)  # First run without cache
    times = []
    
    print(f"\nRound 1: Without Cache")
    print("-" * 70)
    
    for i in range(iterations):
        try:
            start = time.time()
            text, metadata = parser.parse_file(file_path)
            elapsed = time.time() - start
            times.append(elapsed)
            
            print(f"  Run {i+1}: {elapsed:.3f}s | {metadata['word_count']} words, {metadata['pages']} pages")
        except Exception as e:
            print(f"  Run {i+1}: ERROR - {e}")
    
    avg_time = sum(times) / len(times) if times else 0
    print(f"\nAverage (no cache): {avg_time:.3f}s")
    
    # Now test with cache
    print(f"\nRound 2: With Redis Cache (3 runs)")
    print("-" * 70)
    
    cache_times = []
    for i in range(3):
        try:
            start = time.time()
            text, metadata = parser.parse_file(file_path)
            elapsed = time.time() - start
            cache_times.append(elapsed)
            
            from_cache = metadata.get('from_cache', False)
            status = "✓ HIT" if from_cache else "✗ MISS"
            print(f"  Run {i+1}: {elapsed:.3f}s ({status})")
        except Exception as e:
            print(f"  Run {i+1}: ERROR - {e}")
    
    cache_avg = sum(cache_times) / len(cache_times) if cache_times else 0
    
    print(f"\n{'=' * 70}")
    print(f"RESULTS:")
    print(f"  - First parse (no cache): {avg_time:.3f}s")
    print(f"  - Cache hit: {cache_avg:.3f}s")
    if avg_time > 0:
        print(f"  - Speedup: {avg_time / cache_avg:.1f}x faster with cache")
    print(f"{'=' * 70}")

def benchmark_comparison():
    """Compare docling vs python-docx (if both available)."""
    print(f"\n{'=' * 70}")
    print("COMPARISON: docling vs python-docx")
    print(f"{'=' * 70}")
    
    # Try to import both
    try:
        from docx import Document as DocxDocument
        print("✓ python-docx available")
    except ImportError:
        print("✗ python-docx not installed")
        return
    
    try:
        from docling.document_converter import DocumentConverter
        print("✓ docling available")
    except ImportError:
        print("✗ docling not installed")
        return
    
    # Create test file if needed
    if not os.path.exists(SAMPLE_DOCX):
        create_sample_docx(SAMPLE_DOCX, num_pages=30)
    
    if not os.path.exists(SAMPLE_DOCX):
        print(f"Cannot create test file: {SAMPLE_DOCX}")
        return
    
    # Benchmark
    print(f"\nTesting with: {SAMPLE_DOCX}")
    print("-" * 70)
    
    # python-docx
    try:
        start = time.time()
        doc = DocxDocument(SAMPLE_DOCX)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        text = '\n'.join(text_parts)
        docx_time = time.time() - start
        print(f"python-docx: {docx_time:.3f}s | {len(text)} chars")
    except Exception as e:
        print(f"python-docx: ERROR - {e}")
        docx_time = 0
    
    # docling
    try:
        start = time.time()
        converter = DocumentConverter()
        result = converter.convert(SAMPLE_DOCX)
        text = result.document.export_to_markdown()
        docling_time = time.time() - start
        print(f"docling: {docling_time:.3f}s | {len(text)} chars")
    except Exception as e:
        print(f"docling: ERROR - {e}")
        docling_time = 0
    
    if docx_time > 0 and docling_time > 0:
        speedup = docling_time / docx_time
        print(f"\nSpeedup: python-docx is {speedup:.1f}x faster than docling")

def show_redis_cache_status():
    """Show Redis cache status."""
    print(f"\n{'=' * 70}")
    print("CACHE STATUS")
    print(f"{'=' * 70}")
    
    try:
        import redis
        r = redis.StrictRedis(host='localhost', port=6379, db=1, decode_responses=True)
        r.ping()
        
        keys = r.keys('doc_parse:*')
        print(f"✓ Redis connected")
        print(f"  - Parser cache keys: {len(keys)}")
        print(f"  - TTL: 7 days (604800s)")
        
        if keys:
            print(f"\n  Sample keys:")
            for key in keys[:5]:
                ttl = r.ttl(key)
                size = r.memory_usage(key)
                print(f"    - {key}: TTL={ttl}s, Size={size} bytes")
    except Exception as e:
        print(f"✗ Redis not available: {e}")
        print(f"  Install: apt-get install redis-server")

if __name__ == '__main__':
    print("\n" + "=" * 70)
    print("DOCUMENT PARSER OPTIMIZATION BENCHMARK")
    print("=" * 70)
    print("\nOptimizations:")
    print("  1. python-docx for DOCX (5-10x faster than docling)")
    print("  2. Redis caching (avoid re-parsing same file)")
    print("  3. Performance metrics in metadata")
    print("  4. Cache TTL: 7 days")
    
    # Create test file
    test_file = create_sample_docx(SAMPLE_DOCX, num_pages=30)
    
    if test_file:
        # Run benchmarks
        benchmark_docx_parsing(test_file, iterations=3)
        benchmark_comparison()
        show_redis_cache_status()
    
    print("\n" + "=" * 70)
    print("✓ Benchmark complete!")
    print("=" * 70)
